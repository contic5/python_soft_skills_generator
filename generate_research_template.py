from docx import Document
from docx.shared import Inches

list_number=0
def add_list_style(list_paragraphs,num_id_list_number_new):
    for i in range(len(list_paragraphs)):
        numPr = list_paragraphs[i]._element.pPr._add_numPr()
        numPr._add_numId().val = num_id_list_number_new

def generate_new_list_style(document):
    #prepare the numberings to have a new numbering, which points to the same abstract numbering, 
    #the style "List Number" also points to but has start override set
    styles = document.styles
    #get numId to which style 'List Number' links
    num_id_list_number = -1
    for style in styles:
        if (style.name == 'List Number'):
            num_id_list_number = style._element.pPr.numPr.numId.val
    #add new numbering linking to same abstractNumId but has startOverride 
    #and get new numId
    num_id_list_number_new = -1
    if (num_id_list_number > -1):        
        ct_numbering = document.part.numbering_part.numbering_definitions._numbering
        ct_num = ct_numbering.num_having_numId(num_id_list_number)
        abstractNumId = ct_num.abstractNumId.val
        ct_num = ct_numbering.add_num(abstractNumId)
        num_id_list_number_new = ct_num.numId
        startOverride = ct_num.add_lvlOverride(0)._add_startOverride()
        startOverride.val = 1
    return num_id_list_number_new

class ResearchTemplate:
    def __init__(self,level,topic,total_pictures,total_videos,questions,folder_name):
        self.level=level
        self.topic=topic
        self.total_pictures=total_pictures
        self.total_videos=total_videos
        self.questions=questions
        self.folder_name=folder_name

    def generate_research_documents(self):
        #Initial Setup
        
        research_document = Document()
        research_document.add_heading(f'{self.topic} - {self.level}', 0)
        research_document.add_heading(f'Today you are going to do some research on {self.topic}!', 1)
        self.write_research(research_document)
        self.write_pictures(research_document)
        self.write_videos(research_document)
        
        research_document.save(f'{self.folder_name}/1 - {self.topic} {self.level} Research Template.docx')

class LLResearchTemplate(ResearchTemplate):
    def write_research(self,research_document):
        #Part 1 Research
        research_document.add_heading(f'Part 1 - Research',2)
        research_document.add_paragraph('Please fill in the chart below using google to help you find the answers. Attach the link to the website you used')
        research_document.add_paragraph('Make sure to put information in your own words. Do not copy and paste information directly.')

        table = research_document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Question'
        hdr_cells[1].text = 'Answer'
        hdr_cells[2].text = 'Website Link'
        for i in range(len(self.questions)):
            row_cells = table.add_row().cells
            row_cells[0].text = self.questions[i]
    
    def write_pictures(self,research_document):
        #Part 2 Pictures
        research_document.add_heading(f'Part 2 - Pictures',2)
        p = research_document.add_paragraph(f'Research and find {self.total_pictures} pictures for {self.topic}. In google images type “{self.topic} Tips.” Click on the picture and copy and paste it into the table below. Attach the link to the website you used below.')
        table = research_document.add_table(rows=1+self.total_pictures, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Pictures'
        hdr_cells[1].text = 'Website'
    
    def write_videos(self,research_document):
        #Part 3 Videos
        research_document.add_heading(f'Part 3 - Videos',2)
        p = research_document.add_paragraph(f'Please watch and find {self.total_videos} YouTube videos about {self.topic} tips for students. These should be videos YOU find fun and interesting. Copy and paste the links below. In the comments column write 1-2 sentences about your favorite part of the video.')

        table = research_document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Video Links'
        hdr_cells[1].text = 'Comments'
        for i in range(self.total_videos):
            row_cells = table.add_row().cells
            row_cells[0].text = f"{i+1}."

class MLResearchTemplate(ResearchTemplate):
    def write_research(self,research_document):
        global list_number

        research_document.add_heading(f'Part 1 - Research',2)
        research_document.add_paragraph("INSERT TABLE HERE")
        num_id_list_number_new=generate_new_list_style(research_document)
        list_paragraphs=[]

        list_paragraphs.append(research_document.add_paragraph("Click above where it says 'INSERT TABLE HERE'. Click Insert, then click on table.",style="List Number"))
        list_paragraphs.append(research_document.add_paragraph("Use your mouse to highlight a 3x6 area.",style="List Number"))
        list_paragraphs.append(research_document.add_paragraph("Once you have a 3x6 click on the box, a table should pop up in your document.",style="List Number"))
        list_paragraphs.append(research_document.add_paragraph("Label the first box in column 1 Question",style="List Number"))
        list_paragraphs.append(research_document.add_paragraph("Label the first box in column 2 Answer",style="List Number"))
        list_paragraphs.append(research_document.add_paragraph("Label the first box in column 3 Website Link",style="List Number"))       
        for i in range(len(self.questions)):
            list_paragraphs.append(research_document.add_paragraph(f"In column 1, row {i+2}, type the question {self.questions[i]}",style="List Number"))
        #list_number+=1
        
        add_list_style(list_paragraphs,num_id_list_number_new)


        research_document.add_heading(f'Part 2 - Research',2)
        research_document.add_paragraph("Please fill in the chart you made above using google to help you find the answers. Type each question into google and find an answer. Type your answers in the box next to the question. Do not copy and paste your answers directly. Then, copy and paste the link into the third box. Attach the link to the website you used to find the answer.")

    def write_pictures(self,research_document):
        global list_number
        num_id_list_number_new=generate_new_list_style(research_document)
        list_paragraphs=[]
        
        research_document.add_heading(f'Part 3 - Pictures',2)
        research_document.add_paragraph("INSERT TABLE HERE")
        list_paragraphs.append(research_document.add_paragraph("Click above where it says 'INSERT TABLE HERE'. Click Insert, then click on table.",style="List Number"))
        list_paragraphs.append(research_document.add_paragraph("Once you have a 2x6 click on the box, a table should pop up in your document.",style="List Number"))
        list_paragraphs.append(research_document.add_paragraph("Label the first box in column 1 Pictures",style="List Number"))
        list_paragraphs.append(research_document.add_paragraph("Label the first box in column 2 Website",style="List Number"))
        #list_number+=1

        add_list_style(list_paragraphs,num_id_list_number_new)

        research_document.add_heading(f'Part 4 - Pictures',2)
        research_document.add_paragraph(f"Research and find 5 pictures for {self.topic}. In google images type {self.topic} Click on the picture and copy and paste it into the table you made above. Attach the link to the website you used in the chart above.")

    def write_videos(self,research_document):
        num_id_list_number_new=generate_new_list_style(research_document)
        list_paragraphs=[]
    
        research_document.add_heading(f'Part 5 - Videos',2)
        research_document.add_paragraph("INSERT TABLE HERE")
        list_paragraphs.append(research_document.add_paragraph("Click above where it says 'INSERT TABLE HERE'. Click Insert, then click on table.",style="List Number"))
        list_paragraphs.append(research_document.add_paragraph("Use your mouse to highlight a 2x6 area.",style="List Number"))
        list_paragraphs.append(research_document.add_paragraph("Once you have a 2x6 click on the box, a table should pop up in your document.",style="List Number"))
        list_paragraphs.append(research_document.add_paragraph("Label the first box in column 1 Video Links",style="List Number"))
        list_paragraphs.append(research_document.add_paragraph("Label the first box in column 2 Comments",style="List Number"))
        #list_number+=1

        add_list_style(list_paragraphs,num_id_list_number_new)

        research_document.add_heading(f'Part 6 - Videos',2)
        research_document.add_paragraph(f"Please find and watch 3 YouTube videos about {self.topic} tips for students. These should be videos YOU find fun and interesting. Copy and paste the links in your table above. When you have finished watching a video write 1-2 sentences about what you liked best about the video in the box next to where you put the link.")

class HLResearchTemplate(ResearchTemplate):
    def write_research(self,research_document):
        global list_number
        num_id_list_number_new=generate_new_list_style(research_document)
        list_paragraphs=[]

        list_paragraphs.append(research_document.add_paragraph("Create a table with 3 columns and 6 rows.",style="List Number"))
        list_paragraphs.append(research_document.add_paragraph("Label the first box in column 1 Question",style="List Number"))
        list_paragraphs.append(research_document.add_paragraph("Label the first box in column 2 Answer",style="List Number"))
        list_paragraphs.append(research_document.add_paragraph("Label the first box in column 3 Website Link",style="List Number"))
        for i in range(len(self.questions)):
           list_paragraphs.append(research_document.add_paragraph(f"In column 1, row {i+2}, type the question {self.questions[i]}",style="List Number"))
        #list_number+=1

        add_list_style(list_paragraphs,num_id_list_number_new)

        research_document.add_heading(f'Part 2 - Research',2)
        research_document.add_paragraph("Please fill in the chart you made. Use Google to help you find the answers. Type all your answers. Do not copy and paste directly. Attach the link to the website you used. You may copy and paste the link.")

    
   
    def write_pictures(self,research_document):
        global list_number
        num_id_list_number_new=generate_new_list_style(research_document)
        list_paragraphs=[]

        list_paragraphs.append(research_document.add_paragraph("Create a table with 2 columns and 6 rows.",style="List Number"))
        list_paragraphs.append(research_document.add_paragraph("Label the first box in column 1 Pictures",style="List Number"))
        list_paragraphs.append(research_document.add_paragraph("Label the first box in column 2 Website",style="List Number"))
        #list_number+=1

        add_list_style(list_paragraphs,num_id_list_number_new)

        research_document.add_heading(f'Part 4 - Pictures',2)
        research_document.add_paragraph(f"Research and find 5 pictures for {self.topic}. In google images type {self.topic} Click on the picture and copy and paste it into the table you made above. Attach the link to the website you used in the chart above.")

    def write_videos(self,research_document):
        global list_number
        num_id_list_number_new=generate_new_list_style(research_document)
        list_paragraphs=[]

        list_paragraphs.append(research_document.add_paragraph("Create a table with 2 columns and 6 rows.",style="List Number"))
        list_paragraphs.append(research_document.add_paragraph("Label the first box in column 1 Video Links",style="List Number"))
        list_paragraphs.append(research_document.add_paragraph("Label the first box in column 2 Comments",style="List Number"))
        #list_number+=1

        add_list_style(list_paragraphs,num_id_list_number_new)

        research_document.add_heading(f'Part 6 - Videos',2)
        research_document.add_paragraph(f"Please find and watch 3 YouTube videos about {self.topic} tips for students. These should be videos YOU find fun and interesting. Copy and paste the links in your table above. When you have finished watching a video write 1-2 sentences about what you liked best about the video in the box next to where you put the link.")

