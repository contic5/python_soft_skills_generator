from docx import Document
from docx.shared import Inches

letters="ABCDEFGHIJKLMNOPQRSTUVWXYZ"
list_number=0

def add_list_style(list_paragraphs,num_id_list_number_new):
    for i in range(len(list_paragraphs)):
        numPr = list_paragraphs[i]._element.pPr._add_numPr()
        numPr._add_numId().val = num_id_list_number_new

def generate_new_list_style(document):
    #prepare the numberings to have a new numbering, which points to the same abstract numbering, 
    #the style "List Number" also points to but has start override set
    styles = document.styles
    #get numId to which style 'List Number' links
    num_id_list_number = -1
    for style in styles:
        if (style.name == 'List Number'):
            num_id_list_number = style._element.pPr.numPr.numId.val
    #add new numbering linking to same abstractNumId but has startOverride 
    #and get new numId
    num_id_list_number_new = -1
    if (num_id_list_number > -1):        
        ct_numbering = document.part.numbering_part.numbering_definitions._numbering
        ct_num = ct_numbering.num_having_numId(num_id_list_number)
        abstractNumId = ct_num.abstractNumId.val
        ct_num = ct_numbering.add_num(abstractNumId)
        num_id_list_number_new = ct_num.numId
        startOverride = ct_num.add_lvlOverride(0)._add_startOverride()
        startOverride.val = 1
    return num_id_list_number_new

class OutlineTemplate:
    def __init__(self,level,topic,total_pictures,total_videos,questions,folder_name):
        self.level=level
        self.topic=topic
        self.total_pictures=total_pictures
        self.total_videos=total_videos
        self.questions=questions
        self.folder_name=folder_name
        print(f"Folder name:{self.folder_name}")

    def generate_outline_documents(self):
        outline_document = Document()
        outline_document.add_heading(f'{self.topic} - {self.level}', 0)

        #Todo steps for preparing outline_document
        outline_document.add_paragraph("To Do: Make double spacing, add bold and delete this paragraph.")
        
        self.write_part1(outline_document)
        self.write_slide_steps(outline_document)
        self.write_website_and_image_steps(outline_document)
        self.write_fillin(outline_document)
        
        outline_document.save(f'{self.folder_name}/2 - {self.topic} {self.level} Outline Template.docx')

    def write_part1(self,outline_document):
        global list_number
        num_id_list_number_new=generate_new_list_style(outline_document)
        list_paragraphs=[]

        #Part 1 Steps
        outline_document.add_heading('Part 1', 2)

        list_paragraphs.append(outline_document.add_paragraph('Go to the Outline Document you worked on yesterday. Before opening it, click the three dots next to the name of the outline_document. Click on Open In Sharepoint. This should bring you to the Internet. Open the outline_document in this tab.', style='List Number'))
        list_paragraphs.append(outline_document.add_paragraph('Go back to Teams but do not close this internet tab completely.', style='List Number'))
        list_paragraphs.append(outline_document.add_paragraph(f'In the outline_document find the word Title at the top. Delete this word and put the name of the title of your outline_document there instead. This title should be creative and include the words “{self.topic}.” ', style='List Number'))
        list_paragraphs.append(outline_document.add_paragraph('Underneath the title, delete the words: Your Name and type your name there instead.',style=f"List Number"))
        list_number+=1

        add_list_style(list_paragraphs,num_id_list_number_new)
        

    def write_slide_steps(self,outline_document):
        #Steps for Question Slides
        global list_number
        for i in range(len(self.questions)):
            
            outline_document.add_heading(f'Part {i+2}', 2)
            num_id_list_number_new=generate_new_list_style(outline_document)
            list_paragraphs=[]

            if i>=2 and self.level=="ML":
                list_paragraphs.append(outline_document.add_paragraph(f"Repeat the previous steps for Slide {i+2}",style="List Number"))
                add_list_style(list_paragraphs,num_id_list_number_new)
                continue

            list_paragraphs.append(outline_document.add_paragraph(f'Find Slide {i+2} in Bold Blue Print',style="List Number"))

            list_paragraphs.append(outline_document.add_paragraph(f'Next to A. Delete the word Question that is already there. Type question {i+1} from part 2 of the Outline Document you worked on yesterday. This should say {self.questions[i]}',style="List Number"))
            list_paragraphs.append(outline_document.add_paragraph(f'Next to to B. Delete the word Answer that is already there. Type the answer to part 1 from the Outline Document you worked on yesterday.',style="List Number"))
            list_paragraphs.append(outline_document.add_paragraph(f'Next to C. Delete the word Picture that is already there. Enter a picture from part 2 of the Outline Document you worked on yesterday.',style="List Number"))
            list_number+=1
            add_list_style(list_paragraphs,num_id_list_number_new)

        #Steps for Video Slide
        num_id_list_number_new=generate_new_list_style(outline_document)
        list_paragraphs=[]
        outline_document.add_heading(f'Part 7', 2)

        list_paragraphs.append(outline_document.add_paragraph('Find Slide 7 in Bold Blue Print.',style="List Number"))
        list_paragraphs.append(outline_document.add_paragraph('Next to B. Delete the word Video that is already there. Choose the best video from part 3 from the Outline Document you worked on yesterday. Copy and paste the link.',style="List Number"))
        list_number+=1
        add_list_style(list_paragraphs,num_id_list_number_new)

        #Steps for Requesting Questions slide
        num_id_list_number_new=generate_new_list_style(outline_document)
        list_paragraphs=[]
        outline_document.add_heading(f'Part 8', 2)

        list_paragraphs.append(outline_document.add_paragraph('Find Slide 8 in Bold Blue Print. ',style="List Number"))
        list_paragraphs.append(outline_document.add_paragraph('Next to B. Delete the word Picture that is already there. Enter a picture from part 2 of the Outline Document you worked on yesterday.',style="List Number"))
        list_number+=1
        add_list_style(list_paragraphs,num_id_list_number_new)


    def write_website_and_image_steps(self,outline_document):
        global list_number

        #Steps for References Slide
        num_id_list_number_new=generate_new_list_style(outline_document)
        list_paragraphs=[]
        outline_document.add_heading(f'Part 9', 2)
        list_paragraphs.append(outline_document.add_paragraph('Find Slide 9 in Bold Blue Print',style="List Number"))
        list_number+=1
        add_list_style(list_paragraphs,num_id_list_number_new)


        #Website Steps
        loc=1
        num_id_list_number_new=generate_new_list_style(outline_document)
        list_paragraphs=[]
        for i in range(len(self.questions)):
            if self.level=="LL" or i<2:
                list_paragraphs.append(outline_document.add_paragraph(f'Next to {letters[loc]}. Delete the words Website for answer on slide {i+2}. Enter the website you used for the answer on slide {i+2}.',style="List Number"))
            else:
                list_paragraphs.append(outline_document.add_paragraph(f'Repeat the previous steps for Slide {i+2}.',style="List Number"))
            loc+=1
        list_number+=1

        #Image Steps
        for i in range(len(self.questions)):
            if self.level=="LL" or i<2:
                list_paragraphs.append(outline_document.add_paragraph(f'Next to {letters[loc]}. Delete the words Website for picture on slide {i+2}. Enter the website you used for the answer on slide {i+2}.',style="List Number"))
            else:
                list_paragraphs.append(outline_document.add_paragraph(f'Repeat the previous steps for Slide {i+2}.',style="List Number"))
            loc+=1

        list_paragraphs.append(outline_document.add_paragraph(f'Next to {letters[loc]}. Delete the words Website for video on slide 7. Enter the website you used for the video on slide 7.',style="List Number"))
        list_number+=1
        add_list_style(list_paragraphs,num_id_list_number_new)

        loc+=1
        outline_document.add_page_break()

    def write_fillin(self,outline_document):
        global list_number
        #Title
        outline_document.add_heading("Title",1)
        outline_document.add_heading("By: Your Name",2)

        #Questions
        for i in range(len(self.questions)):
            num_id_list_number_new=generate_new_list_style(outline_document)
            list_paragraphs=[]
            outline_document.add_heading(f"Slide {i+2}",2)
            list_paragraphs.append(outline_document.add_paragraph("Question",style="List Number"))
            list_paragraphs.append(outline_document.add_paragraph("Answer",style="List Number"))
            list_paragraphs.append(outline_document.add_paragraph("Picture",style="List Number"))
            list_number+=1  
            add_list_style(list_paragraphs,num_id_list_number_new)

            if self.level=="ML" and i>=1:
                return

        #Video Slide
        num_id_list_number_new=generate_new_list_style(outline_document)
        list_paragraphs=[]
        outline_document.add_heading(f"Slide 7",2)
        list_paragraphs.append(outline_document.add_paragraph(f"{self.topic} Video",style="List Number"))
        list_paragraphs.append(outline_document.add_paragraph(f"Video Link",style="List Number"))
        list_number+=1
        add_list_style(list_paragraphs,num_id_list_number_new)

        #Request Questions Slide
        num_id_list_number_new=generate_new_list_style(outline_document)
        list_paragraphs=[]
        outline_document.add_heading(f"Slide 8",2)
        list_paragraphs.append(outline_document.add_paragraph(f"Questions",style="List Number"))
        list_paragraphs.append(outline_document.add_paragraph(f"Picture",style="List Number"))
        list_number+=1
        add_list_style(list_paragraphs,num_id_list_number_new)

        #References Slide
        num_id_list_number_new=generate_new_list_style(outline_document)
        list_paragraphs=[]
        outline_document.add_heading(f"Slide 9 (References Slide)",2)
        list_paragraphs.append(outline_document.add_paragraph(f"References",style="List Number"))
        for i in range(len(self.questions)):
            list_paragraphs.append(outline_document.add_paragraph(f"Website for Answer on Slide {i+2}",style="List Number"))
        for i in range(len(self.questions)):
            list_paragraphs.append(outline_document.add_paragraph(f"Picture for Answer on Slide {i+2}",style="List Number"))
        
        list_paragraphs.append(outline_document.add_paragraph(f"Website for Video on Slide 7",style="List Number"))
        add_list_style(list_paragraphs,num_id_list_number_new)

class LLOutlineTemplate(OutlineTemplate):
    pass

class MLOutlineTemplate(OutlineTemplate):
    pass

class HLOutlineTemplate(OutlineTemplate):
    def write_highlevel(self,outline_document):
        outline_document.add_heading(f"Part 1",2)
        num_id_list_number_new=generate_new_list_style(outline_document)
        list_paragraphs=[]
        list_paragraphs.append(outline_document.add_paragraph(f"Write the title first. This should contain the words {self.topic}",style="List Number"))
        list_paragraphs.append(outline_document.add_paragraph("Underneath the title, put the words Presenter: Your Name",style="List Number"))
        list_paragraphs.append(outline_document.add_paragraph("Press enter twice. Then, on the left side of the page number you are going to begin to create your outline. ",style="List Number"))
        list_paragraphs.append(outline_document.add_paragraph("Type Slide 2 in Bold Print. Press enter.",style="List Number"))
        list_paragraphs.append(outline_document.add_paragraph("On the next line you are going to indent using the TAB button and type A.",style="List Number"))
        list_paragraphs.append(outline_document.add_paragraph(f"Next to A. type the first question from part 2 of the document you worked on yesterday. This Should say {self.questions[0]}. Then press enter",style="List Number"))
        list_paragraphs.append(outline_document.add_paragraph("Below the A. type B. And type the answer to part 2. Then press enter.  ",style="List Number"))
        list_paragraphs.append(outline_document.add_paragraph("Below the B. type C. And enter the picture from part 4 of the document you worked on yesterday. Then press enter.",style="List Number"))
        list_paragraphs.append(outline_document.add_paragraph("Backspace so that your cursor is aligned with the text that says Slide 2. ",style="List Number"))
        list_paragraphs.append(outline_document.add_paragraph("Follow steps 6-12 for the rest of questions/slides.",style="List Number"))
        add_list_style(list_paragraphs,num_id_list_number_new)

        num_id_list_number_new=generate_new_list_style(outline_document)
        list_paragraphs=[]
        outline_document.add_paragraph("The last three slides should be as follows.")
        list_paragraphs.append(outline_document.add_paragraph("A video slide containing the video of your choice ",style="List Number"))
        list_paragraphs.append(outline_document.add_paragraph("A “questions” slide - o	This slide give your audience the opportunity to ask you any questions they have. You do not need to include your own questions on here",style="List Number"))
        list_paragraphs.append(outline_document.add_paragraph("A “References” slide - o	This lists all the websites you got your answers from. ",style="List Number"))
        add_list_style(list_paragraphs,num_id_list_number_new)

        outline_document.add_paragraph("Congratulations! You have finished your outline and are now ready to move on to the PowerPoint section. ")
    
    def generate_outline_documents(self):
        outline_document = Document()
        outline_document.add_heading(f'{self.topic} - {self.level}', 0)

        #Todo steps for preparing outline_document
        outline_document.add_paragraph("To Do: Make double spacing, add bold and delete this paragraph.")
        
        self.write_highlevel(outline_document)
        outline_document.save(f'{self.folder_name}/2 - {self.topic} {self.level} Outline Template.docx')